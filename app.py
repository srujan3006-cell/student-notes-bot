import streamlit as st
from docx import Document
import docx
from io import BytesIO
import google.generativeai as genai

# Configure your API Key securely!
genai.configure(api_key=st.secrets["GEMINI_API_KEY"]) 
model = genai.GenerativeModel('gemini-2.5-flash') 

def generate_docx(text, title):
    doc = Document()
    doc.add_heading(title, 0)
    
    in_code_block = False
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 1. Handle Code Blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
            
        if in_code_block:
            p = doc.add_paragraph(lines[i])
            p.runs[0].font.name = 'Courier New'
            i += 1
            continue

        # 2. Handle Tables (Lines starting with |)
        if line.startswith('|') and i + 1 < len(lines) and '---' in lines[i+1]:
            # Identify all lines that belong to this table
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                line_content = lines[i].strip()
                
                # Filter out the Markdown separator line (e.g., |---|---|)
                # We check if the line contains ONLY dashes, pipes, and spaces
                is_separator = all(c in '|- :' for c in line_content)
                
                if not is_separator:
                    # Clean the line and split by |
                    row = [cell.strip() for cell in line_content.split('|') if cell.strip()]
                    if row: # Only add if the row isn't empty
                        table_data.append(row)
                i += 1
            
            if table_data:
                # Create Table in Word
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_data):
                    for c_idx, cell_value in enumerate(row_data):
                        table.cell(r_idx, c_idx).text = cell_value
            continue

        # 3. Handle Headings
        if line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=1)
        # 4. Handle Bullets
        elif line.startswith('*') or line.startswith('-'):
            doc.add_paragraph(line.replace('*', '').replace('-', '').strip(), style='List Bullet')
        # 5. Handle Regular Text
        elif line != "":
            doc.add_paragraph(line)
            
        i += 1
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Sidebar UI ---
st.sidebar.title("⚙️ Document Settings")
template = st.sidebar.selectbox(
    "Choose a Template:", 
    ["Standard Notes", "Engineering Lab Report", "Exam Study Guide"]
)

# --- Main Streamlit UI ---
st.title("🎓 Student Note Architect")
st.write("Generate structured project notes in Word format instantly.")

topic = st.text_input("What subject do you need notes for?")
depth = st.select_slider("Select detail level:", options=["Summary", "Detailed", "Comprehensive"])

if st.button("Generate Document"):
    if not topic:
        st.warning("Please enter a topic first!")
    else:
        with st.spinner(f"Writing your {template.lower()}..."):
            
            # --- Style Prompting ---
            if template == "Engineering Lab Report":
                style_instruction = "Structure this as a formal lab report. Include: Objective, Components/Bill of Materials, Circuit Architecture setup, Code/Implementation logic, and Conclusion."
            elif template == "Exam Study Guide":
                style_instruction = "Structure this as an exam review. Include: Key definitions, important formulas, and likely exam questions."
            else:
                style_instruction = "Structure this as standard academic notes with a logical flow of concepts."
            
           # --- Clean, Text-Only Prompt ---
            prompt = f"""
            Write {depth} notes on {topic}. {style_instruction} 
            Use '##' for section headers and bullet points for key facts. 
            If there are formulas, you MUST provide them in a Markdown table with two columns: 'Formula' and 'Function/Description'.
            Do not use conversational filler.
            """
            
            try:
                # Generate Text
                response = model.generate_content(prompt)
                
                # Create the docx in memory (This line used to crash if response.text was empty)
                docx_file = generate_docx(response.text, topic)
                
                st.success(f"{template} generated successfully!")
                st.download_button(
                    label="Download .docx File",
                    data=docx_file,
                    file_name=f"{topic.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except ValueError:
                # Catch the empty response error gracefully
                st.error("⚠️ The AI safety filters blocked this response, or it returned empty text. Please try tweaking your topic slightly!")
            except Exception as e:
                # Catch any other random errors
                st.error(f"⚠️ An unexpected error occurred: {e}")
